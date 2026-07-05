HTML_TEMPLATE = '<!DOCTYPE html>\n<html lang="en">\n<head>\n<meta charset="UTF-8">\n<script src="qrc:///qtwebchannel/qwebchannel.js"></script>\n<style>\n@font-face {\n  font-family: \'OpenDyslexic\';\n  font-weight: normal;\n  src: url(\'data:font/ttf;base64,__FONT__\') format(\'truetype\');\n}\n@font-face {\n  font-family: \'OpenDyslexic\';\n  font-weight: bold;\n  src: url(\'data:font/ttf;base64,__FONT_BOLD__\') format(\'truetype\');\n}\n*,*::before,*::after{box-sizing:border-box;margin:0;padding:0;}\nhtml,body{height:100%;background:transparent;}\nbody{font-family:system-ui,sans-serif;background:transparent;color:#c8960c;display:flex;flex-direction:column;height:100%;}\n#drop-hint{flex:1;display:flex;flex-direction:column;align-items:center;justify-content:center;gap:12px;color:#9a7218;font-size:13px;cursor:default;user-select:none;}\n#drop-hint .icon{font-size:3rem;opacity:0.5;}\n#drop-hint strong{color:#c8960c;font-size:14px;}\n#drop-hint small{font-size:10px;color:#6a5210;}\n#status{padding:8px 16px;font-size:11px;color:#9a7218;text-align:center;display:none;background:rgba(15,15,18,0.9);border-top:1px solid rgba(200,150,12,0.2);}\n#status.error{color:#f87171;display:block;}\n#status.loading{display:block;}\n.progress-bar{width:100%;height:3px;background:rgba(200,150,12,0.1);margin-top:5px;border-radius:2px;overflow:hidden;}\n.progress-fill{height:100%;background:#c8960c;width:0%;transition:width 0.3s;}\n#reader-wrap{flex:1;display:none;flex-direction:column;overflow:hidden;}\n#reader{flex:1;overflow-y:auto;padding:32px 44px;font-family:\'OpenDyslexic\',sans-serif;font-size:17px;line-height:1.8;letter-spacing:0.05em;color:#c8960c;background:transparent;}\n#reader h1,#reader h2,#reader h3,#reader h4,#reader h5,#reader h6{font-family:\'OpenDyslexic\',sans-serif;color:#e8b830;margin:0.9em 0 0.35em;}\n#reader h1{font-size:1.7em;}#reader h2{font-size:1.35em;}#reader h3{font-size:1.1em;}\n#reader p{margin-bottom:0.85em;}\n#reader pre,#reader code{font-family:\'OpenDyslexic\',monospace;background:rgba(200,150,12,0.07);border:1px solid rgba(200,150,12,0.2);border-radius:4px;padding:0.2em 0.4em;font-size:0.87em;color:#e8b830;}\n#reader pre{padding:0.9em;overflow-x:auto;display:block;}\n#reader pre code{background:none;border:none;padding:0;}\n#reader blockquote{border-left:3px solid #510E8C;padding-left:1em;margin-left:0;color:#9a7218;}\n#reader table{border-collapse:collapse;width:100%;margin-bottom:1em;font-size:0.87em;}\n#reader th,#reader td{border:1px solid rgba(200,150,12,0.3);padding:5px 10px;text-align:left;}\n#reader th{background:rgba(200,150,12,0.1);color:#e8b830;}\n#reader ul,#reader ol{padding-left:1.5em;margin-bottom:0.85em;}\n#reader li{margin-bottom:0.2em;}\n#reader a{color:#8f6fff;}\n#reader img{max-width:100%;border-radius:4px;}\n#reader hr{border:none;border-top:1px solid rgba(200,150,12,0.2);margin:1.2em 0;}\n.image-panel{display:flex;gap:14px;flex-wrap:wrap;}\n.panel-box{flex:1;min-width:240px;}\n.panel-box h3{font-size:10px;color:#9a7218;margin-bottom:5px;text-transform:uppercase;letter-spacing:0.1em;font-family:system-ui;}\n.panel-box canvas,.panel-box img{max-width:100%;border-radius:4px;display:block;border:1px solid rgba(200,150,12,0.2);}\n.pdf-page{margin-bottom:14px;border-radius:4px;overflow:hidden;border:1px solid rgba(200,150,12,0.15);}\n.pdf-page .pg-label{background:rgba(200,150,12,0.08);padding:2px 10px;font-size:10px;color:#9a7218;font-family:system-ui;}\n.pdf-page .pg-text{padding:16px 22px;}\n#translate-bar{display:none;padding:8px 12px;background:rgba(15,15,18,0.95);border-top:1px solid rgba(200,150,12,0.15);gap:8px;align-items:center;flex-shrink:0;}\n#translate-bar select{background:#0a0a0e;color:#c8960c;border:1px solid #c8960c;border-radius:5px;font-size:11px;padding:3px 6px;}\n#translate-bar button{background:#0d0d12;color:#c8960c;border:1px solid #c8960c;border-radius:5px;font-size:11px;padding:4px 12px;cursor:pointer;}\n#translate-bar button:hover{border-color:#e8b830;color:#e8b830;}\n#translate-bar .engine-tag{font-size:10px;color:#9a7218;}\n::-webkit-scrollbar{width:5px;}::-webkit-scrollbar-track{background:transparent;}\n::-webkit-scrollbar-thumb{background:rgba(200,150,12,0.5);border-radius:3px;}\n#img-viewer{display:flex;flex-direction:column;width:100%;}\n#overlay-view{display:flex;flex-direction:column;align-items:stretch;gap:10px;width:100%;}\n#text-view{display:flex;flex-direction:column;gap:10px;width:100%;}\n.img-toggle-btn{background:#0d0d12;color:#c8960c;border:1px solid #c8960c;border-radius:5px;font-size:11px;padding:5px 14px;cursor:pointer;align-self:flex-start;margin-top:8px;}\n.img-toggle-btn:hover{border-color:#e8b830;color:#e8b830;}\n.text-view-actions{display:flex;gap:8px;}\n.img-action-btn{background:#0d0d12;color:#c8960c;border:1px solid #c8960c;border-radius:5px;font-size:11px;padding:5px 14px;cursor:pointer;}\n.img-action-btn:hover{border-color:#e8b830;color:#e8b830;}\n.ocr-text-panel{white-space:pre-wrap;word-break:break-word;font-family:\'OpenDyslexic\',sans-serif;font-size:0.88em;color:#c8960c;background:rgba(200,150,12,0.05);border:1px solid rgba(200,150,12,0.2);border-radius:4px;padding:16px;margin:0;width:100%;box-sizing:border-box;letter-spacing:normal;line-height:1.6;}\n.tts-w.tts-active{background:rgba(200,150,12,0.35);border-radius:3px;}\n</style>\n</head>\n<body>\n<div id="drop-hint">\n  <span class="icon">&#x2750;</span>\n  <strong>Open a file</strong>\n  <div style="color:#9a7218;font-size:11px">Use the toolbar button, drag a file onto the window, or paste (Ctrl+V)</div>\n  <small>TXT · MD · HTML · CSV · PDF · DOCX · RTF · JSON · PNG · JPG · GIF · BMP · WEBP</small>\n</div>\n<div id="status"></div>\n<div id="reader-wrap">\n  <div id="reader"></div>\n  <div id="translate-bar">\n    <span style="font-size:11px;color:#9a7218;">Translate to:</span>\n    <select id="lang-select">\n      <option value="en" selected>English</option>\n      <option value="es">Spanish</option>\n      <option value="fr">French</option>\n      <option value="de">German</option>\n      <option value="it">Italian</option>\n      <option value="pt">Portuguese</option>\n      <option value="nl">Dutch</option>\n      <option value="pl">Polish</option>\n      <option value="ru">Russian</option>\n      <option value="zh-CN">Chinese (Simplified)</option>\n      <option value="ja">Japanese</option>\n      <option value="ko">Korean</option>\n      <option value="ar">Arabic</option>\n      <option value="hi">Hindi</option>\n      <option value="sv">Swedish</option>\n      <option value="no">Norwegian</option>\n      <option value="da">Danish</option>\n      <option value="fi">Finnish</option>\n      <option value="tr">Turkish</option>\n      <option value="uk">Ukrainian</option>\n    </select>\n    <button id="translate-btn">Translate</button>\n    <span class="engine-tag" id="engine-tag"></span>\n  </div>\n</div>\n\n<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js" crossorigin="anonymous"></script>\n<script src="https://cdnjs.cloudflare.com/ajax/libs/mammoth/1.6.0/mammoth.browser.min.js" crossorigin="anonymous"></script>\n<script src="https://cdnjs.cloudflare.com/ajax/libs/marked/9.1.6/marked.min.js" crossorigin="anonymous"></script>\n<script>\npdfjsLib.GlobalWorkerOptions.workerSrc=\'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js\';\n\nlet pyBridge=null;\nnew QWebChannel(qt.webChannelTransport,ch=>{\n  pyBridge=ch.objects.bridge;\n  const _ce=console.error.bind(console),_cw=console.warn.bind(console),_ci=console.log.bind(console);\n  console.error=(...a)=>{_ce(...a);try{pyBridge.log_js(a.map(String).join(\' \'));}catch(_){}};\n  console.warn=(...a)=>{_cw(...a);try{pyBridge.log_js_warn(a.map(String).join(\' \'));}catch(_){}};\n  console.log=(...a)=>{_ci(...a);try{pyBridge.log_js_info(a.map(String).join(\' \'));}catch(_){}};\n  window.onerror=(msg,src,line,col,err)=>{\n    try{pyBridge.log_js(`${msg} (${src}:${line}:${col})${err?\' — \'+err.stack:\'\'}`);}catch(_){}\n    return false;\n  };\n  window.addEventListener(\'unhandledrejection\',e=>{\n    try{pyBridge.log_js(\'Unhandled rejection: \'+(e.reason?.stack||e.reason||e));}catch(_){}\n  });\n});\n\n// ── State ─────────────────────────────────────────────────────────────────────\nlet _currentText = \'\';\nlet _lastImgBlob=null;\nlet _lastOcrData=null;\nlet _lastImgElement=null;\nlet _themeBg=null;\nlet _themeFg=null;\nlet _bionicCanvas=false;\nwindow.setThemeBg=(bg,fg)=>{_themeBg=bg;_themeFg=fg;if(_lastOcrData&&_lastImgElement)_redrawCanvas(_lastImgElement,_lastOcrData);};\n\nfunction setStatus(msg,type=\'loading\',pct=null){\n  const el=document.getElementById(\'status\');\n  if(!msg){el.style.display=\'none\';el.className=\'\';el.innerHTML=\'\';return;}\n  el.className=type;el.style.display=\'block\';\n  el.innerHTML=msg+(pct!=null?`<div class="progress-bar"><div class="progress-fill" style="width:${pct}%"></div></div>`:\'\');\n}\nfunction esc(t){return String(t).replace(/&/g,\'&amp;\').replace(/</g,\'&lt;\').replace(/>/g,\'&gt;\').replace(/"/g,\'&quot;\');}\nfunction show(){\n  document.getElementById(\'drop-hint\').style.display=\'none\';\n  document.getElementById(\'reader-wrap\').style.display=\'flex\';\n  document.getElementById(\'translate-bar\').style.display=\'flex\';\n}\nfunction reset(){\n  document.getElementById(\'reader-wrap\').style.display=\'none\';\n  document.getElementById(\'drop-hint\').style.display=\'flex\';\n  document.getElementById(\'translate-bar\').style.display=\'none\';\n  document.getElementById(\'reader\').innerHTML=\'\';\n  document.getElementById(\'engine-tag\').textContent=\'\';\n  _currentText=\'\';\n  setStatus(\'\');\n}\nfunction setReaderText(html, plainText){\n  const rdr=document.getElementById(\'reader\');if(!rdr)return;\n  rdr.innerHTML=html;\n  _currentText=plainText||rdr.innerText;\n  if(typeof _wrapTtsSpans===\'function\')_wrapTtsSpans();\n}\n\nwindow.loadFileFromPython=(name,ext,b64)=>receiveFileData(JSON.stringify({name,ext,b64}));\nwindow.resetView=reset;\n\n// ── Translate ─────────────────────────────────────────────────────────────────\ndocument.getElementById(\'translate-btn\').addEventListener(\'click\', async ()=>{\n  const lang=document.getElementById(\'lang-select\').value;\n  const text=_currentText.trim();\n  if(!text){setStatus(\'Nothing to translate\',\'error\');return;}\n  setStatus(\'Translating…\',\'loading\',50);\n  console.log(`translate: lang=${lang}, chars=${text.length}`);\n  pyBridge.translate_text(text, lang, res=>{\n    setStatus(\'\');\n    try{\n      const d=JSON.parse(res);\n      if(d.error){setStatus(\'Translation error: \'+d.error,\'error\');return;}\n      document.getElementById(\'reader\').innerHTML=\'<pre style="white-space:pre-wrap;word-break:break-word">\'+esc(d.result)+\'</pre>\';\n      _currentText=d.result;\n      document.getElementById(\'engine-tag\').textContent=\'via \'+d.engine;\n      console.log(\'translate: done via \'+d.engine);\n    }catch(e){setStatus(\'Parse error: \'+e.message,\'error\');}\n  });\n});\n\n// ── File receive ──────────────────────────────────────────────────────────────\nasync function receiveFileData(json){\n  if(!json)return;\n  let d;try{d=JSON.parse(json);}catch{return;}\n  if(d.error){setStatus(d.error,\'error\');return;}\n  const{name,ext,b64}=d;\n  show();\n  try{\n    const bytes=Uint8Array.from(atob(b64),c=>c.charCodeAt(0));\n    const blob=new Blob([bytes]);\n    if([\'png\',\'jpg\',\'jpeg\',\'gif\',\'bmp\',\'webp\'].includes(ext)){\n      await handleImage(blob);\n    }else if(ext===\'pdf\'){\n      await handlePDF(bytes.buffer);\n    }else if(ext===\'docx\'){\n      await handleDOCX(bytes.buffer);\n    }else if(ext===\'md\'||ext===\'markdown\'){\n      handleMD(new TextDecoder().decode(bytes));\n    }else if(ext===\'html\'||ext===\'htm\'){\n      handleHTML(new TextDecoder().decode(bytes));\n    }else if(ext===\'csv\'){\n      handleCSV(new TextDecoder().decode(bytes));\n    }else if(ext===\'rtf\'){\n      handleRTF(new TextDecoder().decode(bytes));\n    }else{\n      handleText(new TextDecoder().decode(bytes));\n    }\n  }catch(e){setStatus(\'Error: \'+e.message,\'error\');console.error(e.stack||e);}\n}\n\nconst reader=document.getElementById(\'reader\');\n\nfunction handleText(t){setStatus(\'\');setReaderText(\'<pre style="white-space:pre-wrap;word-break:break-word">\'+esc(t)+\'</pre>\',t);}\nfunction handleMD(t){setStatus(\'\');setReaderText(marked.parse(t),t);}\nfunction handleHTML(t){\n  setStatus(\'\');\n  const doc=new DOMParser().parseFromString(t,\'text/html\');\n  doc.querySelectorAll(\'script,style,link\').forEach(el=>el.remove());\n  const html=doc.body?doc.body.innerHTML:t;\n  setReaderText(html,doc.body?doc.body.innerText:t);\n  reader.querySelectorAll(\'*\').forEach(el=>{el.style.fontFamily=\'\';el.style.fontSize=\'\';el.style.color=\'\';el.style.background=\'\';});\n}\nfunction handleCSV(t){\n  setStatus(\'\');\n  const rows=t.trim().split(\'\\n\').map(r=>{\n    const cells=[];let cur=\'\',inQ=false;\n    for(const c of r){if(c===\'"\'){inQ=!inQ;}else if(c===\',\'&&!inQ){cells.push(cur.trim());cur=\'\';}else cur+=c;}\n    cells.push(cur.trim());return cells;\n  });\n  let h=\'<div style="overflow-x:auto"><table><thead><tr>\';\n  rows[0].forEach(c=>h+=`<th>${esc(c)}</th>`);\n  h+=\'</tr></thead><tbody>\';\n  rows.slice(1).forEach(r=>{h+=\'<tr>\';r.forEach(c=>h+=`<td>${esc(c)}</td>`);h+=\'</tr>\';});\n  const html=h+\'</tbody></table></div>\';\n  setReaderText(html,t);\n}\nfunction handleRTF(t){\n  setStatus(\'\');\n  let plain=t.replace(/\\\\[a-z]+-?[0-9]* ?/g,\'\').replace(/[{}\\\\]/g,\'\').replace(/[^\\x20-\\x7E\\n\\r]/g,\'\').replace(/ +/g,\' \').trim();\n  setReaderText(\'<pre style="white-space:pre-wrap;word-break:break-word">\'+esc(plain)+\'</pre>\',plain);\n}\nasync function handleDOCX(buf){\n  setStatus(\'Converting DOCX…\');\n  const r=await mammoth.convertToHtml({arrayBuffer:buf});\n  setStatus(\'\');\n  reader.innerHTML=r.value;\n  reader.querySelectorAll(\'*\').forEach(el=>{el.style.fontFamily=\'\';el.style.fontSize=\'\';el.style.color=\'\';el.style.background=\'\';});\n  _currentText=reader.innerText;\n}\nasync function handlePDF(buf){\n  setStatus(\'Loading PDF…\',\'loading\',5);\n  const pdf=await pdfjsLib.getDocument({data:buf}).promise;\n  const total=pdf.numPages;\n  reader.innerHTML=\'\';\n  let allText=\'\';\n  for(let p=1;p<=total;p++){\n    setStatus(`Page ${p} / ${total}…`,\'loading\',10+(p/total)*85);\n    const page=await pdf.getPage(p);\n    const vp=page.getViewport({scale:1.5});\n    const canvas=document.createElement(\'canvas\');\n    canvas.width=vp.width;canvas.height=vp.height;\n    await page.render({canvasContext:canvas.getContext(\'2d\'),viewport:vp}).promise;\n    const tc=await page.getTextContent();\n    const txt=tc.items.map(i=>i.str).join(\' \');\n    allText+=txt+\'\\n\\n\';\n    const div=document.createElement(\'div\');\n    div.className=\'pdf-page\';\n    const pgLbl=document.createElement(\'div\');pgLbl.className=\'pg-label\';pgLbl.textContent=\'Page \'+p;\n    canvas.style.cssText=\'max-width:100%;display:block;border-radius:4px;\';\n    const pgTxt=document.createElement(\'div\');pgTxt.className=\'pg-text\';pgTxt.textContent=txt;\n    div.appendChild(pgLbl);div.appendChild(canvas);div.appendChild(pgTxt);\n    reader.appendChild(div);\n  }\n  _currentText=allText;\n  setStatus(\'\');\n}\nfunction _redrawCanvas(img,ocr){\n  const validWords=(ocr.words||[]).filter(w=>w.text.trim()&&(w.bbox.x1-w.bbox.x0)>0&&(w.bbox.y1-w.bbox.y0)>0);\n  if(!validWords.length)return;\n  const slot=document.getElementById(\'canvas-slot\');\n  if(!slot)return;\n  const canvas=slot.querySelector(\'canvas\');\n  if(!canvas)return;\n  const ctx=canvas.getContext(\'2d\');\n  ctx.drawImage(img,0,0);\n  const pad=6;\n  const ux0=Math.max(0,Math.min(...validWords.map(w=>w.bbox.x0))-pad);\n  const uy0=Math.max(0,Math.min(...validWords.map(w=>w.bbox.y0))-pad);\n  const ux1=Math.min(canvas.width,Math.max(...validWords.map(w=>w.bbox.x1))+pad);\n  const uy1=Math.min(canvas.height,Math.max(...validWords.map(w=>w.bbox.y1))+pad);\n  const uw=ux1-ux0,uh=uy1-uy0;\n  const sy=Math.max(0,uy0-10),sw=Math.min(uw,canvas.width-ux0);\n  const bgData=ctx.getImageData(ux0,sy,sw,1).data;\n  let br=0,bg_=0,bb=0;\n  for(let i=0;i<sw;i++){br+=bgData[i*4];bg_+=bgData[i*4+1];bb+=bgData[i*4+2];}\n  br=Math.round(br/sw);bg_=Math.round(bg_/sw);bb=Math.round(bb/sw);\n  const sampledBg=`rgb(${br},${bg_},${bb})`;\n  const bgColor=_themeBg||sampledBg;\n  const lum=0.299*br+0.587*bg_+0.114*bb;\n  const textColor=_themeFg||(lum>140?\'#1a0f00\':\'#c8960c\');\n  ctx.fillStyle=bgColor;\n  ctx.fillRect(0,0,canvas.width,canvas.height);\n  const heights=validWords.map(w=>w.bbox.y1-w.bbox.y0).sort((a,b)=>a-b);\n  const medH=heights[Math.floor(heights.length/2)];\n  const globalFs=Math.max(8,Math.round(medH*0.72));\n  ctx.textBaseline=\'top\';\n  const lineThresh=medH*0.6;\n  const lines=[];\n  const sorted=[...validWords].sort((a,b)=>a.bbox.y0-b.bbox.y0);\n  sorted.forEach(word=>{\n    const y=word.bbox.y0;\n    const line=lines.find(l=>Math.abs(l.baseY-y)<lineThresh);\n    if(line){line.words.push(word);line.baseY=Math.round((line.baseY+y)/2);}\n    else lines.push({baseY:y,words:[word]});\n  });\n  lines.forEach(line=>{\n    const lw=line.words.slice().sort((a,b)=>a.bbox.x0-b.bbox.x0);\n    if(!lw.length)return;\n    ctx.font=`${globalFs}px OpenDyslexic,sans-serif`;\n    const lineX0=lw[0].bbox.x0,lineX1=lw[lw.length-1].bbox.x1,lineW=lineX1-lineX0;\n    const textWidths=lw.map(w=>{ctx.font=`${globalFs}px OpenDyslexic,sans-serif`;return Math.min(ctx.measureText(w.text).width,w.bbox.x1-w.bbox.x0);});\n    const totalTextW=textWidths.reduce((a,b)=>a+b,0);\n    const gap=lw.length>1?(lineW-totalTextW)/(lw.length-1):0;\n    const evenGap=Math.min(gap,globalFs*0.6);\n    let cx=lineX0;\n    lw.forEach((word,i)=>{\n      ctx.fillStyle=textColor;\n      if(_bionicCanvas){\n        const half=Math.ceil(word.text.length/2);\n        const b=word.text.slice(0,half),r2=word.text.slice(half);\n        ctx.font=`bold ${globalFs}px OpenDyslexic,sans-serif`;\n        ctx.fillText(b,cx,line.baseY);ctx.fillText(b,cx+0.5,line.baseY);ctx.fillText(b,cx+1,line.baseY);\n        const bw=ctx.measureText(b).width;\n        ctx.font=`${globalFs}px OpenDyslexic,sans-serif`;\n        ctx.fillText(r2,cx+bw,line.baseY);\n      }else{\n        ctx.font=`${globalFs}px OpenDyslexic,sans-serif`;\n        ctx.fillText(word.text,cx,line.baseY);\n      }\n      cx+=textWidths[i]+(i<lw.length-1?evenGap:0);\n    });\n  });\n}\nasync function handleImage(blob){\n  try{\n    _lastImgBlob=blob;\n    console.log(\'handleImage: start, blob size=\'+blob.size);\n    setStatus(\'Running OCR…\',\'loading\',20);\n    const url=URL.createObjectURL(blob);\n    const b64=await new Promise((res,rej)=>{\n      const fr=new FileReader();\n      fr.onload=()=>res(fr.result.split(\',\')[1]);\n      fr.onerror=rej;\n      fr.readAsDataURL(blob);\n    });\n    console.log(\'handleImage: sending to Python OCR\');\n    const ocrJson=await new Promise((res,rej)=>{\n      pyBridge.ocr_image(b64,result=>res(result));\n    });\n    const ocr=JSON.parse(ocrJson);\n    if(ocr.error){console.error(\'OCR error: \'+ocr.error);setStatus(\'OCR error: \'+ocr.error,\'error\');}\n    console.log(\'handleImage: OCR done, words=\'+(ocr.words||[]).length);\n    setStatus(\'Rendering…\',\'loading\',85);\n    const img=new Image();img.src=url;\n    await new Promise((res,rej)=>{img.onload=res;img.onerror=()=>rej(new Error(\'Image load failed\'));});\n    console.log(\'handleImage: image loaded \'+img.width+\'x\'+img.height);\n    _lastImgElement=img;_lastOcrData=ocr;\n    const canvas=document.createElement(\'canvas\');\n    canvas.width=img.width;canvas.height=img.height;\n    const ctx=canvas.getContext(\'2d\');\n    ctx.drawImage(img,0,0);\n    // Filter valid words\n    const validWords=(ocr.words||[]).filter(word=>{\n      if(!word.text.trim())return false;\n      const b=word.bbox,w=b.x1-b.x0,h=b.y1-b.y0;\n      return w>0&&h>0;\n    });\n    if(!validWords.length){setStatus(\'\');return;}\n    // Compute union bbox of all words\n    const pad=6;\n    const ux0=Math.max(0,Math.min(...validWords.map(w=>w.bbox.x0))-pad);\n    const uy0=Math.max(0,Math.min(...validWords.map(w=>w.bbox.y0))-pad);\n    const ux1=Math.min(canvas.width,Math.max(...validWords.map(w=>w.bbox.x1))+pad);\n    const uy1=Math.min(canvas.height,Math.max(...validWords.map(w=>w.bbox.y1))+pad);\n    const uw=ux1-ux0,uh=uy1-uy0;\n    // Sample background from a strip just above the text region\n    const sy=Math.max(0,uy0-10),sw=Math.min(uw,canvas.width-ux0);\n    const bgData=ctx.getImageData(ux0,sy,sw,1).data;\n    let br=0,bg=0,bb=0;\n    for(let i=0;i<sw;i++){br+=bgData[i*4];bg+=bgData[i*4+1];bb+=bgData[i*4+2];}\n    br=Math.round(br/sw);bg=Math.round(bg/sw);bb=Math.round(bb/sw);\n    const sampledBg=`rgb(${br},${bg},${bb})`;\n    const bgColor=_themeBg||sampledBg;\n    // Pick text color: theme fg if set, else luminance-based\n    const lum=0.299*br+0.587*bg+0.114*bb;\n    const textColor=_themeFg||(lum>140?\'#1a0f00\':\'#c8960c\');\n    // Fill entire canvas with theme color, or just text bbox if sampled\n    ctx.fillStyle=bgColor;\n    ctx.fillRect(0,0,canvas.width,canvas.height);\n    // Compute median word height for stable font size\n    const heights=validWords.map(word=>word.bbox.y1-word.bbox.y0).sort((a,b)=>a-b);\n    const medH=heights[Math.floor(heights.length/2)];\n    const globalFs=Math.max(8,Math.round(medH*0.72));\n    ctx.textBaseline=\'top\';\n    ctx.fillStyle=textColor;\n    // Cluster words into lines by y0 proximity, snap each line to shared baseline\n    const lineThresh=medH*0.6;\n    const lines=[];\n    const sorted=[...validWords].sort((a,b)=>a.bbox.y0-b.bbox.y0);\n    sorted.forEach(word=>{\n      const y=word.bbox.y0;\n      const line=lines.find(l=>Math.abs(l.baseY-y)<lineThresh);\n      if(line){line.words.push(word);line.baseY=Math.round((line.baseY+y)/2);}\n      else lines.push({baseY:y,words:[word]});\n    });\n    let wc=0;\n    lines.forEach(line=>{\n      const by=line.baseY;\n      // Sort words left to right\n      const lw=line.words.slice().sort((a,b)=>a.bbox.x0-b.bbox.x0);\n      if(!lw.length)return;\n      // Measure total OpenDyslexic text width for this line\n      ctx.font=`${globalFs}px OpenDyslexic,sans-serif`;\n      const lineX0=lw[0].bbox.x0;\n      const lineX1=lw[lw.length-1].bbox.x1;\n      const lineW=lineX1-lineX0;\n      const textWidths=lw.map(word=>{\n        ctx.font=`${globalFs}px OpenDyslexic,sans-serif`;\n        return Math.min(ctx.measureText(word.text).width,word.bbox.x1-word.bbox.x0);\n      });\n      const totalTextW=textWidths.reduce((a,b)=>a+b,0);\n      // Even gap between words (de-justify)\n      const gap=lw.length>1?(lineW-totalTextW)/(lw.length-1):0;\n      const evenGap=Math.min(gap,globalFs*0.6);// cap gap at 60% of font size\n      let cx=lineX0;\n      lw.forEach((word,i)=>{\n        ctx.fillStyle=textColor;\n        if(_bionicCanvas){\n          const half=Math.ceil(word.text.length/2);\n          const b=word.text.slice(0,half),r2=word.text.slice(half);\n          ctx.font=`bold ${globalFs}px OpenDyslexic,sans-serif`;\n          ctx.fillText(b,cx,by);ctx.fillText(b,cx+0.5,by);ctx.fillText(b,cx+1,by);\n          const bw=ctx.measureText(b).width;\n          ctx.font=`${globalFs}px OpenDyslexic,sans-serif`;\n          ctx.fillText(r2,cx+bw,by);\n        }else{\n          ctx.font=`${globalFs}px OpenDyslexic,sans-serif`;\n          ctx.fillText(word.text,cx,by);\n        }\n        cx+=textWidths[i]+(i<lw.length-1?evenGap:0);\n        wc++;\n      });\n    });\n    console.log(\'handleImage: painted \'+wc+\' words\');\n    canvas.style.cssText=\'max-width:100%;width:100%;display:block;border-radius:4px;border:1px solid rgba(200,150,12,0.2)\';\n    const ocrText=ocr.text||\'\';\n    reader.innerHTML=\n      \'<div id="img-viewer">\'\n      +\'<div id="overlay-view">\'\n      +\'<button class="img-toggle-btn" id="to-text-btn" style="margin:4px 0 2px 0;">Show Extracted Text ▸</button>\'\n      +\'<div id="canvas-slot" style="width:100%;display:flex;flex-direction:column;"></div>\'\n      +\'</div>\'\n      +\'<div id="text-view" style="display:none">\'\n      +\'<div class="text-view-actions">\'\n      +\'<button class="img-action-btn" id="copy-btn">⎘ Copy to Clipboard</button>\'\n      +\'<button class="img-action-btn" id="save-btn">❐ Save As…</button>\'\n      +\'</div>\'\n      +\'<pre id="ocr-text-pre" class="ocr-text-panel"></pre>\'\n      +\'<button class="img-toggle-btn" id="to-overlay-btn">◂ Show Overlay</button>\'\n      +\'</div>\'\n      +\'</div>\';\n    document.getElementById(\'canvas-slot\').appendChild(canvas);\n    document.getElementById(\'ocr-text-pre\').textContent=ocrText;\n    document.getElementById(\'to-text-btn\').onclick=()=>{\n      document.getElementById(\'overlay-view\').style.display=\'none\';\n      document.getElementById(\'text-view\').style.display=\'flex\';\n    };\n    document.getElementById(\'to-overlay-btn\').onclick=()=>{\n      document.getElementById(\'text-view\').style.display=\'none\';\n      document.getElementById(\'overlay-view\').style.display=\'flex\';\n    };\n    document.getElementById(\'copy-btn\').onclick=()=>{\n      const btn=document.getElementById(\'copy-btn\');\n      pyBridge.copy_text(ocrText,res=>{\n        btn.textContent=\'✓ Copied!\';\n        setTimeout(()=>btn.textContent=\'⎘ Copy to Clipboard\',2000);\n      });\n    };\n    document.getElementById(\'save-btn\').onclick=()=>{\n      pyBridge.save_text(ocrText,res=>{\n        try{\n          const d=JSON.parse(res);\n          if(d.error)setStatus(\'Save error: \'+d.error,\'error\');\n          else if(d.saved)setStatus(\'Saved: \'+d.saved,\'\');\n        }catch(e){}\n      });\n    };\n    _currentText=ocrText;\n    setStatus(\'\');\n    console.log(\'handleImage: done\');\n  }catch(e){\n    console.error(\'handleImage failed: \'+(e.stack||e));\n    setStatus(\'Image error: \'+e.message,\'error\');\n  }\n}\n// ── Bionic Reading ───────────────────────────────────────────────────────────\nlet _bionic=false;\nfunction _applyBionic(el){\n  const walker=document.createTreeWalker(el,NodeFilter.SHOW_TEXT,{acceptNode:n=>n.parentNode.nodeName!==\'SCRIPT\'&&n.parentNode.nodeName!==\'STYLE\'?NodeFilter.FILTER_ACCEPT:NodeFilter.FILTER_REJECT});\n  const nodes=[];while(walker.nextNode())nodes.push(walker.currentNode);\n  nodes.forEach(node=>{\n    const text=node.nodeValue;if(!text.trim())return;\n    const frag=document.createDocumentFragment();\n    text.split(/(\\s+)/).forEach(part=>{\n      if(/\\s/.test(part)){frag.appendChild(document.createTextNode(part));return;}\n      if(!part)return;\n      const half=Math.ceil(part.length/2);\n      const b=document.createElement(\'b\');b.className=\'bio\';b.style.fontWeight=\'900\';b.textContent=part.slice(0,half);\n      frag.appendChild(b);frag.appendChild(document.createTextNode(part.slice(half)));\n    });\n    node.parentNode.replaceChild(frag,node);\n  });\n}\nfunction _removeBionic(el){\n  el.querySelectorAll(\'.bio\').forEach(b=>{\n    const next=b.nextSibling;\n    const combined=b.textContent+(next&&next.nodeType===3?next.nodeValue:\'\');\n    const tn=document.createTextNode(combined);\n    if(next&&next.nodeType===3)b.parentNode.removeChild(next);\n    b.parentNode.replaceChild(tn,b);\n  });\n}\nwindow.toggleBionic=()=>{_bionic=!_bionic;_bionicCanvas=_bionic;const r=document.getElementById(\'reader\');if(_bionic)_applyBionic(r);else _removeBionic(r);if(_lastImgElement&&_lastOcrData)_redrawCanvas(_lastImgElement,_lastOcrData);};\n// ── Reading Ruler ─────────────────────────────────────────────────────────────\nlet _ruler=false;\nconst rulerEl=document.createElement(\'div\');rulerEl.id=\'ruler\';\nrulerEl.style.cssText=\'position:fixed;left:0;right:0;height:3.5em;background:rgba(200,150,12,0.15);pointer-events:none;display:none;z-index:999;transition:top 0.05s;\';\ndocument.body.appendChild(rulerEl);\nconst _magCnv=document.createElement(\'canvas\');\n_magCnv.style.cssText=\'position:fixed;pointer-events:none;display:none;z-index:1000;border-top:1px solid rgba(200,150,12,0.5);border-bottom:1px solid rgba(200,150,12,0.5);\';\ndocument.body.appendChild(_magCnv);\nlet _magTextEl=null;\ndocument.addEventListener(\'mousemove\',e=>{\n  if(!_ruler)return;\n  const ry=e.clientY-16;\n  rulerEl.style.top=ry+\'px\';\n  const ic=document.querySelector(\'#canvas-slot canvas\');\n  if(ic&&ic.offsetParent!==null){\n    const r=ic.getBoundingClientRect();\n    const rh=parseFloat(getComputedStyle(rulerEl).height)||28;\n    const sy=ic.height/r.height;\n    const srcH=(rh/1.35)*sy;\n    const srcY=Math.max(0,Math.min((e.clientY-r.top-rh/2)*sy,ic.height-srcH));\n    _magCnv.width=Math.round(r.width);_magCnv.height=Math.round(rh);\n    _magCnv.style.left=r.left+\'px\';_magCnv.style.top=ry+\'px\';\n    _magCnv.style.width=r.width+\'px\';_magCnv.style.height=rh+\'px\';_magCnv.style.display=\'block\';\n    const mx=_magCnv.getContext(\'2d\');\n    mx.clearRect(0,0,_magCnv.width,_magCnv.height);\n    mx.drawImage(ic,0,srcY,ic.width,srcH,0,0,_magCnv.width,_magCnv.height);\n    if(_magTextEl){_magTextEl.style.transform=\'\';_magTextEl.style.zIndex=\'\';_magTextEl.style.position=\'\';_magTextEl=null;}\n  }else{\n    _magCnv.style.display=\'none\';\n    const hit=document.elementFromPoint(e.clientX,e.clientY);\n    let tgt=hit;while(tgt&&tgt!==document.body){const ds=getComputedStyle(tgt).display;if((ds===\'block\'||ds===\'flex\')&&tgt.offsetHeight<100)break;tgt=tgt.parentElement;}\n    if(tgt!==_magTextEl){\n      if(_magTextEl){_magTextEl.style.transform=\'\';_magTextEl.style.zIndex=\'\';_magTextEl.style.position=\'\';}\n      if(tgt&&tgt!==document.body&&tgt!==document.documentElement){tgt.style.transform=\'scale(1.14)\';tgt.style.transformOrigin=\'left center\';tgt.style.position=\'relative\';tgt.style.zIndex=\'500\';}\n      _magTextEl=(tgt&&tgt!==document.body)?tgt:null;\n    }\n  }\n});\nwindow.toggleRuler=(on)=>{_ruler=on;rulerEl.style.display=on?\'block\':\'none\';if(!on){_magCnv.style.display=\'none\';if(_magTextEl){_magTextEl.style.transform=\'\';_magTextEl.style.zIndex=\'\';_magTextEl.style.position=\'\';_magTextEl=null;}}};\n// ── TTS ───────────────────────────────────────────────────────────────────────\nconst synth=window.speechSynthesis;let _ttsUtterance=null;let _ttsWordMap=[];\nfunction _buildTtsWordMap(){\n  _ttsWordMap=[];const text=_currentText;const regex=/\\S+/g;let m;\n  while((m=regex.exec(text))!==null)_ttsWordMap.push({start:m.index,end:m.index+m[0].length});\n}\nwindow.ttsPlay=()=>{\n  if(synth.speaking)synth.cancel();const text=_currentText.trim();if(!text)return;\n  _buildTtsWordMap();\n  _ttsUtterance=new SpeechSynthesisUtterance(text);_ttsUtterance.rate=0.9;\n  _ttsUtterance.onboundary=(e)=>{\n    if(e.name!==\'word\')return;\n    document.querySelectorAll(\'.tts-w.tts-active\').forEach(el=>el.classList.remove(\'tts-active\'));\n    const ci=e.charIndex;const match=_ttsWordMap.find(w=>w.start<=ci&&ci<w.end);\n    if(match){document.querySelectorAll(\'.tts-w\').forEach(sp=>{if(parseInt(sp.dataset.start)===match.start){sp.classList.add(\'tts-active\');sp.scrollIntoView({block:\'nearest\'});}});}\n  };\n  _ttsUtterance.onend=()=>{document.querySelectorAll(\'.tts-w.tts-active\').forEach(el=>el.classList.remove(\'tts-active\'));};\n  synth.speak(_ttsUtterance);\n};\nwindow.ttsPause=()=>{synth.paused?synth.resume():synth.pause();};\nwindow.ttsStop=()=>{synth.cancel();document.querySelectorAll(\'.tts-w.tts-active\').forEach(el=>el.classList.remove(\'tts-active\'));};\nfunction _wrapTtsSpans(){\n  const r=document.getElementById(\'reader\');if(!r)return;\n  const fullText=_currentText;let searchFrom=0;\n  const walker=document.createTreeWalker(r,NodeFilter.SHOW_TEXT,{acceptNode:n=>!n.parentNode.classList.contains(\'tts-w\')&&n.parentNode.nodeName!==\'SCRIPT\'?NodeFilter.FILTER_ACCEPT:NodeFilter.FILTER_REJECT});\n  const nodes=[];while(walker.nextNode())nodes.push(walker.currentNode);\n  nodes.forEach(node=>{\n    const text=node.nodeValue;const frag=document.createDocumentFragment();\n    const regex=/(\\S+)|(\\s+)/g;let m;\n    while((m=regex.exec(text))!==null){\n      if(m[2]!==undefined){frag.appendChild(document.createTextNode(m[2]));}\n      else{const word=m[1];const idx=fullText.indexOf(word,searchFrom);const sp=document.createElement(\'span\');sp.className=\'tts-w\';sp.dataset.start=idx>=0?idx:searchFrom;sp.dataset.end=idx>=0?idx+word.length:searchFrom+word.length;sp.textContent=word;frag.appendChild(sp);if(idx>=0)searchFrom=idx+word.length;}\n    }\n    node.parentNode.replaceChild(frag,node);\n  });\n}\n\n\n</script>\n</body>\n</html>'

#!/usr/bin/env python3
"""
DyslexEye
pip install PyQt6 PyQt6-WebEngine deep-translator pytesseract pillow
python dyslexic_reader.py
"""
import sys, base64, json, math, time, logging, traceback
from pathlib import Path
from datetime import datetime
from io import BytesIO

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QFileDialog, QLabel, QPushButton, QComboBox, QSlider, QFrame, QSizePolicy,
    QSizeGrip, QTextEdit, QLineEdit
)
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWebEngineCore import QWebEngineSettings
from PyQt6.QtWebChannel import QWebChannel
from PyQt6.QtCore import (
    QUrl, pyqtSlot, pyqtSignal, QObject, Qt, QTimer, QPointF,
    QByteArray, QBuffer, QIODeviceBase, QThread
)
from PyQt6.QtGui import (
    QColor, QPainter, QLinearGradient, QRadialGradient,
    QPalette, QAction, QKeySequence, QDragEnterEvent, QDropEvent,
    QFont, QFontDatabase, QCursor
)

# ── Logging ───────────────────────────────────────────────────────────────────
def _init_log() -> logging.Logger:
    root = Path(__file__).parent
    log_dir = root / "logs"
    log_dir.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%A-%B-%d-%H_%M_%S")
    log_file = log_dir / f"{stamp}.txt"
    fmt = logging.Formatter("[%(asctime)s] %(levelname)s — %(message)s", datefmt="%H:%M:%S")
    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setFormatter(fmt)
    log = logging.getLogger("reader")
    log.setLevel(logging.DEBUG)
    log.addHandler(fh)
    # Redirect stdout/stderr to log file so nothing goes to the console
    sys.stdout = open(log_file, "a", encoding="utf-8", buffering=1)
    sys.stderr = sys.stdout
    return log

log = _init_log()

def _except_hook(exc_type, exc_value, exc_tb):
    log.error("Unhandled exception:\n" + "".join(traceback.format_exception(exc_type, exc_value, exc_tb)))
    sys.__excepthook__(exc_type, exc_value, exc_tb)
sys.excepthook = _except_hook

# ── Font ──────────────────────────────────────────────────────────────────────
_FONT_PATH = Path(__file__).parent / "OpenDyslexic3-Regular.ttf"
_FONT_BOLD_PATH = Path(__file__).parent / "OpenDyslexic3-Bold.ttf"
if _FONT_PATH.exists():
    _FONT_B64 = base64.b64encode(_FONT_PATH.read_bytes()).decode()
    log.info(f"Font loaded: {_FONT_PATH}")
else:
    _FONT_B64 = ""
    log.warning(f"Font not found at {_FONT_PATH}")
if _FONT_BOLD_PATH.exists():
    _FONT_BOLD_B64 = base64.b64encode(_FONT_BOLD_PATH.read_bytes()).decode()
    log.info(f"Bold font loaded: {_FONT_BOLD_PATH}")
else:
    _FONT_BOLD_B64 = ""
    log.warning(f"Bold font not found at {_FONT_BOLD_PATH}")

HTML_CONTENT = HTML_TEMPLATE.replace("__FONT__", _FONT_B64).replace("__FONT_BOLD__", _FONT_BOLD_B64)


# ── Tesseract path helper ─────────────────────────────────────────────────────
def _find_tesseract():
    import shutil, os
    if shutil.which("tesseract"):
        return None
    for candidate in [
        r"C:\Users\jakob\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]:
        if os.path.exists(candidate):
            return candidate
    return None

# ── Qt font family (populated after QApplication is created) ──────────────────
_FONT_FAMILY: str = ""

# ── Palette ───────────────────────────────────────────────────────────────────
_BTN  = ("background:#0d0d12;color:#c8960c;border:1px solid #c8960c;"
         "border-radius:5px;font-size:11px;padding:4px 12px;min-height:24px;")
_BTN_H= ("background:#0d0d12;color:#e8b830;border:1px solid #e8b830;"
         "border-radius:5px;font-size:11px;padding:4px 12px;min-height:24px;")
_LH   = "color:#9a7218;font-size:10px;background:transparent;"
_COMBO= ("background:#0a0a0e;color:#c8960c;border:1px solid #c8960c;"
         "border-radius:5px;font-size:11px;padding:3px 6px;min-height:24px;")

SUPPORTED = (
    "All supported (*.txt *.md *.markdown *.html *.htm *.csv *.pdf "
    "*.docx *.rtf *.json *.log *.py *.js *.ts *.css *.xml "
    "*.png *.jpg *.jpeg *.gif *.bmp *.webp);;"
    "Text (*.txt *.md *.html *.htm *.csv *.rtf *.json *.log *.py *.js);;"
    "PDF (*.pdf);;Word (*.docx);;Images (*.png *.jpg *.jpeg *.gif *.bmp *.webp);;"
    "All (*)"
)

def _encode_file(path: str) -> str:
    p = Path(path)
    log.info(f"Encoding: {p.name} ({p.stat().st_size} bytes)")
    try:
        b64 = base64.b64encode(p.read_bytes()).decode()
        log.debug(f"Encoded OK: {p.name}")
        return json.dumps({"name": p.name, "ext": p.suffix.lstrip(".").lower(), "b64": b64})
    except Exception as e:
        log.error(f"Encode failed {p.name}: {e}")
        return json.dumps({"error": str(e)})

# ── Bridge ────────────────────────────────────────────────────────────────────
class Bridge(QObject):
    def __init__(self, win):
        super().__init__()
        self._win = win

    @pyqtSlot(result=str)
    def open_file(self) -> str:
        log.info("Open file dialog")
        path, _ = QFileDialog.getOpenFileName(self._win, "Open file", str(Path.home()), SUPPORTED)
        if not path:
            log.debug("Dialog cancelled")
            return ""
        log.info(f"Selected: {path}")
        return _encode_file(path)

    @pyqtSlot(str)
    def log_js(self, msg: str): log.error(f"[JS] {msg}")
    @pyqtSlot(str, result=str)
    def fetch_url(self, url: str) -> str:
        import urllib.request, html.parser as _hp, re as _re
        log.info(f"fetch_url: {url}")
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=12) as resp:
                charset = resp.headers.get_content_charset() or "utf-8"
                raw = resp.read().decode(charset, errors="replace")
        except Exception as e:
            log.error(f"fetch_url failed: {e}")
            return json.dumps({"error": str(e)})
        class _Strip(_hp.HTMLParser):
            def __init__(self):
                super().__init__(convert_charrefs=True)
                self._skip = 0; self._buf = []
            def handle_starttag(self, tag, attrs):
                if tag in ("script","style","head","nav","footer","noscript"): self._skip += 1
                if tag in ("p","div","h1","h2","h3","h4","h5","li","br","tr"): self._buf.append("\n\n")
            def handle_endtag(self, tag):
                if tag in ("script","style","head","nav","footer","noscript"): self._skip = max(0, self._skip-1)
            def handle_data(self, data):
                if not self._skip: self._buf.append(data)
        p = _Strip(); p.feed(raw)
        text = _re.sub(r"\n{3,}", "\n\n", "".join(p._buf)).strip()
        return json.dumps({"text": text})

    @pyqtSlot(str)
    def log_js_warn(self, msg: str): log.warning(f"[JS] {msg}")
    @pyqtSlot(str)
    def log_js_info(self, msg: str): log.info(f"[JS] {msg}")

    @pyqtSlot(str, result=str)
    def ocr_image(self, b64: str) -> str:
        log.info("ocr_image: called")
        try:
            import pytesseract
            from PIL import Image
            cmd = _find_tesseract()
            if cmd:
                pytesseract.pytesseract.tesseract_cmd = cmd
                log.info(f"Using Tesseract at {cmd}")
            img_bytes = base64.b64decode(b64)
            img = Image.open(BytesIO(img_bytes)).convert("RGB")
            log.info(f"ocr_image: {img.width}x{img.height}")
            # Auto-invert dark-background images so Tesseract reads light-on-dark correctly
            import numpy as np
            arr = np.array(img)
            avg_lum = float(0.299*arr[:,:,0].mean() + 0.587*arr[:,:,1].mean() + 0.114*arr[:,:,2].mean())
            ocr_img = Image.fromarray(255 - arr) if avg_lum < 128 else img
            log.info(f"ocr_image: avg_lum={avg_lum:.1f}, inverted={avg_lum < 128}")
            data = pytesseract.image_to_data(ocr_img, output_type=pytesseract.Output.DICT)
            words = []
            # Group tokens into lines for clean extracted text
            lines = {}
            for i in range(len(data["text"])):
                txt = data["text"][i].strip()
                conf = int(data["conf"][i]) if data["conf"][i] != -1 else -1
                if txt:
                    words.append({"text": txt, "conf": conf, "bbox": {
                        "x0": data["left"][i], "y0": data["top"][i],
                        "x1": data["left"][i] + data["width"][i],
                        "y1": data["top"][i] + data["height"][i],
                    }})
                # Group by paragraph only — join lines within a para into flowing text
                key = (data["block_num"][i], data["par_num"][i])
                if key not in lines:
                    lines[key] = []
                if txt:
                    lines[key].append(txt)
            import re
            full_text = "\n\n".join(
                " ".join(toks) for toks in lines.values() if toks
            ).strip()
            # Fix missing spaces where sentences were joined without space (e.g. "word.Next")
            full_text = re.sub(r'([a-z\.\]\/])([A-Z])', r'\1 \2', full_text)
            full_text = re.sub(r'([.!?])[^a-zA-Z0-9]*([A-Z])', r'\1 \2', full_text)
            log.info(f"ocr_image: {len(words)} words")
            return json.dumps({"text": full_text, "words": words})
        except ImportError:
            log.warning("pytesseract not installed")
            return json.dumps({"error": "pip install pytesseract pillow", "text": "", "words": []})
        except Exception as e:
            log.error(f"ocr_image failed: {e}\n{traceback.format_exc()}")
            return json.dumps({"error": str(e), "text": "", "words": []})

    @pyqtSlot(str, result=str)
    def copy_text(self, text: str) -> str:
        QApplication.clipboard().setText(text)
        log.info(f"copy_text: {len(text)} chars copied")
        return json.dumps({"ok": True})

    @pyqtSlot(str, result=str)
    def save_text(self, text: str) -> str:
        log.info("save_text: dialog")
        path, _ = QFileDialog.getSaveFileName(
            self._win, "Save extracted text", str(Path.home() / "extracted_text.txt"),
            "Text (*.txt);;CSV (*.csv);;Markdown (*.md);;Word (*.docx);;Rich Text (*.rtf);;All (*)"
        )
        if not path:
            return json.dumps({"saved": None})
        try:
            ext = Path(path).suffix.lower()
            if ext == ".docx":
                try:
                    from docx import Document
                    doc = Document()
                    for line in text.split("\n"):
                        doc.add_paragraph(line)
                    doc.save(path)
                except ImportError:
                    log.warning("python-docx not installed, saving as plain text")
                    Path(path).write_text(text, encoding="utf-8")
            else:
                Path(path).write_text(text, encoding="utf-8")
            log.info(f"save_text: saved to {path}")
            return json.dumps({"saved": Path(path).name})
        except Exception as e:
            log.error(f"save_text failed: {e}\n{traceback.format_exc()}")
            return json.dumps({"error": str(e)})

    @pyqtSlot(str, str, result=str)
    def translate_text(self, text: str, target_lang: str) -> str:
        log.info(f"translate_text: target={target_lang}, chars={len(text)}")
        if not text.strip():
            return json.dumps({"error": "No text", "result": ""})
        try:
            from deep_translator import GoogleTranslator
            result = GoogleTranslator(source="auto", target=target_lang).translate(text)
            log.info(f"translate_text: Google OK, {len(result)} chars")
            return json.dumps({"result": result, "engine": "Google"})
        except Exception as e:
            log.warning(f"translate_text: Google failed ({e}), trying argostranslate")
        try:
            import argostranslate.translate
            installed = argostranslate.translate.get_installed_languages()
            codes = [l.code for l in installed]
            src = "en"
            if src not in codes or target_lang not in codes:
                missing = [c for c in [src, target_lang] if c not in codes]
                msg = f"argostranslate packs missing: {missing}"
                log.warning(msg)
                return json.dumps({"error": msg, "result": ""})
            src_lang = next(l for l in installed if l.code == src)
            tgt_lang = next(l for l in installed if l.code == target_lang)
            result = src_lang.get_translation(tgt_lang).translate(text)
            log.info(f"translate_text: argostranslate OK")
            return json.dumps({"result": result, "engine": "argostranslate"})
        except ImportError:
            msg = "pip install deep-translator"
            log.error(msg)
            return json.dumps({"error": msg, "result": ""})
        except Exception as e:
            log.error(f"translate_text failed: {e}\n{traceback.format_exc()}")
            return json.dumps({"error": str(e), "result": ""})


# ── OCR worker thread ─────────────────────────────────────────────────────────
class OCRWorker(QThread):
    done = pyqtSignal(str)

    def __init__(self, img_bytes: bytes):
        super().__init__()
        self._img_bytes = img_bytes

    def run(self):
        try:
            import pytesseract
            from PIL import Image
            from collections import defaultdict
            cmd = _find_tesseract()
            if cmd:
                pytesseract.pytesseract.tesseract_cmd = cmd
            img = Image.open(BytesIO(self._img_bytes))
            data = pytesseract.image_to_data(
                img, output_type=pytesseract.Output.DICT, config="--psm 3"
            )
            lines: dict = defaultdict(list)
            for i in range(len(data["text"])):
                word = data["text"][i].strip()
                if not word:
                    continue
                try:
                    conf = int(data["conf"][i])
                except (ValueError, TypeError):
                    conf = 0
                if conf < 0:
                    continue
                key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
                lines[key].append(word)
            parts: list = []
            prev_par = None
            for key in sorted(lines.keys()):
                block, par, _ = key
                cur_par = (block, par)
                if prev_par is not None and cur_par != prev_par:
                    parts.append("")
                parts.append(" ".join(lines[key]))
                prev_par = cur_par
            self.done.emit("\n".join(parts).strip())
        except Exception as e:
            log.error(f"OCRWorker: {e}")
            self.done.emit(f"[OCR error: {e}]")


class FileLoader(QThread):
    """Base64-encode a file off the main thread."""
    done = pyqtSignal(str)
    def __init__(self, path):
        super().__init__()
        self._path = path
    def run(self):
        self.done.emit(_encode_file(self._path))


class HTMLFetchWorker(QThread):
    """Fetch raw HTML for layout mode."""
    done = pyqtSignal(str)
    def __init__(self, url):
        super().__init__()
        self._url = url
    def run(self):
        import urllib.request, ssl as _ssl
        try:
            ctx = _ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = _ssl.CERT_NONE
            req = urllib.request.Request(self._url, headers={'User-Agent':'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=15, context=ctx) as resp:
                cs = resp.headers.get_content_charset() or 'utf-8'
                self.done.emit(resp.read().decode(cs, errors='replace'))
        except Exception as e:
            log.error(f'HTMLFetchWorker: {e}')
            self.done.emit(f'<html><body><p style="color:#f87171">Failed to load: {e}</p></body></html>')


class URLWorker(QThread):
    done = pyqtSignal(str)
    def __init__(self, url):
        super().__init__()
        self._url = url
    def run(self):
        import urllib.request, html.parser as _hp, re as _re
        try:
            import ssl as _ssl
            ctx = _ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = _ssl.CERT_NONE
            req = urllib.request.Request(self._url, headers={'User-Agent':'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=15, context=ctx) as resp:
                cs = resp.headers.get_content_charset() or 'utf-8'
                raw = resp.read().decode(cs, errors='replace')
        except Exception as e:
            log.error(f'URLWorker: {e}')
            self.done.emit(json.dumps({'error': str(e)}))
            return
        # 1. Strip scripts, styles, and chrome elements entirely
        for tag in ('script','style','noscript','nav','header','footer',
                    'aside','menu','form','dialog','banner'):
            raw = _re.sub(rf'<{tag}[^>]*>[\s\S]*?</{tag}>', '', raw, flags=_re.IGNORECASE)
        # 2. Remove elements whose id/class smell like chrome
        _CHROME = ('nav','menu','sidebar','header','footer','banner',
                   'cookie','toc','toolbar','search','widget','ad','promo',
                   'breadcrumb','pagination','related','share','social',
                   'comment','masthead','utility','skip','ribbon')
        raw = _re.sub(
            r'<(?:div|section|ul|ol|span|figure|aside)[^>]+(?:id|class)="[^"]*(?:' +
            '|'.join(_CHROME) + r')[^"]*"[^>]*>[\s\S]{0,4000}?</(?:div|section|ul|ol|span|figure|aside)>',
            '', raw, flags=_re.IGNORECASE)
        # 3. Prefer <main> or <article> if present
        for scope_tag in ('main', 'article'):
            m = _re.search(rf'<{scope_tag}[^>]*>([\s\S]*?)</{scope_tag}>', raw, flags=_re.IGNORECASE)
            if m:
                raw = m.group(1)
                break
        # 4. Strip tags and decode entities
        def _striptags(s):
            s = _re.sub(r'<[^>]+>', ' ', s)
            for ent, ch in [('&nbsp;',' '),('&amp;','&'),('&lt;','<'),('&gt;','>'),('&quot;','"')]:
                s = s.replace(ent, ch)
            s = _re.sub(r'&#?\w+;', '', s)
            return ' '.join(s.split())
        # 5. Extract paragraphs and headings
        blocks = _re.findall(r'<(p|h[1-6])[^>]*>([\s\S]*?)</\1>', raw, flags=_re.IGNORECASE)
        parts = [_striptags(b) for _, b in blocks]
        text = '\n\n'.join(p for p in parts if len(p) > 30)
        self.done.emit(json.dumps({'text': text}))

# ── Overlay window ────────────────────────────────────────────────────────────
class OverlayWindow(QWidget):
    _INTERVALS = [500, 1000, 2000, 3000, 5000]

    def __init__(self, main_win):
        super().__init__()
        self._main = main_win
        self._drag_pos = None
        self._worker = None
        self._busy = False

        self.setWindowFlags(
            Qt.WindowType.FramelessWindowHint |
            Qt.WindowType.WindowStaysOnTopHint
        )
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setMinimumSize(320, 180)
        self.resize(540, 340)

        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)

        self._glass = QFrame()
        self._glass.setObjectName("OGlass")
        self._glass.setStyleSheet(
            "#OGlass{background:rgba(10,10,14,218);"
            "border:1px solid rgba(143,111,255,100);border-radius:10px;}"
        )
        gl = QVBoxLayout(self._glass)
        gl.setContentsMargins(6, 4, 6, 6)
        gl.setSpacing(4)

        self._bar = QWidget()
        self._bar.setStyleSheet("background:transparent;")
        self._bar.setFixedHeight(28)
        self._bar.setCursor(Qt.CursorShape.SizeAllCursor)
        bl = QHBoxLayout(self._bar)
        bl.setContentsMargins(4, 0, 4, 0)
        bl.setSpacing(6)

        icon = QLabel("\u2315")
        icon.setStyleSheet("font-size:13px;background:transparent;")
        bl.addWidget(icon)
        title_lbl = QLabel("Overlay Mode")
        title_lbl.setStyleSheet("color:#c8960c;font-size:11px;font-weight:bold;background:transparent;")
        bl.addWidget(title_lbl)
        self._status_dot = QLabel("\u25cf")
        self._status_dot.setStyleSheet("color:#4ade80;font-size:10px;background:transparent;")
        bl.addWidget(self._status_dot)
        bl.addStretch()

        iv_lbl = QLabel("Refresh"); iv_lbl.setStyleSheet(_LH); bl.addWidget(iv_lbl)
        self._iv_combo = QComboBox()
        self._iv_combo.addItems(["0.5s","1s","2s","3s","5s"])
        self._iv_combo.setCurrentIndex(1)
        self._iv_combo.setStyleSheet(_COMBO); self._iv_combo.setFixedWidth(58)
        self._iv_combo.currentIndexChanged.connect(self._update_interval)
        bl.addWidget(self._iv_combo)

        fs_lbl = QLabel("Size"); fs_lbl.setStyleSheet(_LH); bl.addWidget(fs_lbl)
        self._fs_combo = QComboBox()
        self._fs_combo.addItems(["10","12","14","16","18","20"])
        self._fs_combo.setCurrentIndex(2)
        self._fs_combo.setStyleSheet(_COMBO); self._fs_combo.setFixedWidth(48)
        self._fs_combo.currentIndexChanged.connect(self._update_fontsize)
        bl.addWidget(self._fs_combo)

        close_btn = QPushButton("\u2715")
        close_btn.setFixedSize(22, 22)
        close_btn.setStyleSheet(
            "background:#1a0808;color:#f87171;border:1px solid #f87171;"
            "border-radius:4px;font-size:10px;"
        )
        close_btn.clicked.connect(self._close_overlay)
        bl.addWidget(close_btn)
        gl.addWidget(self._bar)

        sep = QFrame(); sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet("background:rgba(200,150,12,0.2);max-height:1px;border:none;")
        gl.addWidget(sep)

        self._text = QTextEdit()
        self._text.setReadOnly(True)
        self._text.setStyleSheet(
            "QTextEdit{background:transparent;color:#c8960c;border:none;}"
            "QScrollBar:vertical{width:5px;background:transparent;}"
            "QScrollBar::handle:vertical{background:rgba(200,150,12,0.5);border-radius:3px;}"
            "QScrollBar::add-line:vertical,QScrollBar::sub-line:vertical{height:0;}"
        )
        self._apply_font(14)
        gl.addWidget(self._text, 1)

        self._grip = QSizeGrip(self)
        self._grip.setStyleSheet("background:transparent;width:14px;height:14px;")
        gl.addWidget(self._grip, 0, Qt.AlignmentFlag.AlignBottom | Qt.AlignmentFlag.AlignRight)
        outer.addWidget(self._glass)

        self._bar.mousePressEvent   = lambda e: self._bar_press(e)
        self._bar.mouseMoveEvent    = lambda e: self._bar_move(e)
        self._bar.mouseReleaseEvent = lambda e: setattr(self, "_drag_pos", None)

        self._timer = QTimer(self)
        self._timer.timeout.connect(self._capture)
        self._timer.start(1000)

    def _apply_font(self, size: int):
        if _FONT_FAMILY:
            self._text.setFont(QFont(_FONT_FAMILY, size))
        else:
            f = self._text.font(); f.setPointSize(size); self._text.setFont(f)

    def _update_interval(self, idx): self._timer.setInterval(self._INTERVALS[idx])
    def _update_fontsize(self, idx): self._apply_font([10,12,14,16,18,20][idx])

    def _bar_press(self, e):
        if e.button() == Qt.MouseButton.LeftButton:
            self._drag_pos = e.globalPosition().toPoint() - self.frameGeometry().topLeft()

    def _bar_move(self, e):
        if self._drag_pos and e.buttons() == Qt.MouseButton.LeftButton:
            self.move(e.globalPosition().toPoint() - self._drag_pos)

    def _close_overlay(self):
        self._timer.stop()
        if self._worker: self._worker.quit()
        self.hide(); self._main.showNormal(); self._main.raise_()

    # ── click-through: Windows ────────────────────────────────────────────────
    def nativeEvent(self, event_type, message):
        try:
            if sys.platform == "win32" and event_type == b"windows_generic_MSG":
                import ctypes
                msg_id = ctypes.c_uint.from_address(int(message) + 8).value
                if msg_id == 0x0084:  # WM_NCHITTEST
                    pos = QCursor.pos()
                    widget = QApplication.widgetAt(pos)
                    if (widget is None or widget is self
                            or widget is self._glass or isinstance(widget, QTextEdit)):
                        return True, -1  # HTTRANSPARENT
                    return True, 1       # HTCLIENT
        except Exception:
            pass
        return super().nativeEvent(event_type, message)

    # ── click-through: Linux X11 ──────────────────────────────────────────────
    def showEvent(self, event):
        super().showEvent(event)
        if sys.platform.startswith("linux"):
            QTimer.singleShot(200, self._x11_update_shape)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if sys.platform.startswith("linux") and self.isVisible():
            QTimer.singleShot(50, self._x11_update_shape)

    def _x11_update_shape(self):
        try:
            import ctypes, ctypes.util
            libx11  = ctypes.CDLL(ctypes.util.find_library("X11")    or "libX11.so.6")
            libxfix = ctypes.CDLL(ctypes.util.find_library("Xfixes") or "libXfixes.so.3")
            dp = libx11.XOpenDisplay(None)
            if not dp: return
            xid = int(self.winId())
            w, h = self.width(), self.height()
            BAR_H, GRIP = 36, 20

            class XRect(ctypes.Structure):
                _fields_ = [("x", ctypes.c_short), ("y", ctypes.c_short),
                             ("width", ctypes.c_ushort), ("height", ctypes.c_ushort)]

            rects = (XRect * 2)(XRect(0, 0, w, BAR_H), XRect(w-GRIP, h-GRIP, GRIP, GRIP))
            region = libxfix.XFixesCreateRegion(dp, rects, 2)
            libxfix.XFixesSetWindowShapeRegion(dp, xid, 2, 0, 0, region)
            libxfix.XFixesDestroyRegion(dp, region)
            libx11.XFlush(dp); libx11.XCloseDisplay(dp)
            log.debug("X11 input shape updated")
        except Exception as e:
            log.warning(f"X11 input shape: {e}")

    # ── capture → OCR ─────────────────────────────────────────────────────────
    def _capture(self):
        if self._busy: return
        self._busy = True
        self._status_dot.setStyleSheet("color:#ffb74d;font-size:10px;background:transparent;")
        geo = self.geometry()
        x, y, w, h = geo.x(), geo.y(), geo.width(), geo.height()
        self.hide()
        QTimer.singleShot(120, lambda: self._do_grab(x, y, w, h))

    def _do_grab(self, x, y, w, h):
        try:
            from PIL import ImageGrab
            img = ImageGrab.grab(bbox=(x, y, x+w, y+h))
            buf = BytesIO(); img.save(buf, 'PNG'); img_bytes = buf.getvalue()
        except Exception as e:
            log.error(f'Overlay grab: {e}')
            self.show(); self._busy = False
            self._status_dot.setStyleSheet('color:#f87171;font-size:10px;background:transparent;')
            return
        self.show()
        self._worker = OCRWorker(img_bytes)
        self._worker.done.connect(self._on_ocr)
        self._worker.start()

    def _on_ocr(self, text: str):
        sv = self._text.verticalScrollBar().value()
        self._text.setPlainText(text)
        self._text.verticalScrollBar().setValue(sv)
        self._status_dot.setStyleSheet("color:#4ade80;font-size:10px;background:transparent;")
        self._worker = None; self._busy = False

# ── Animated background ───────────────────────────────────────────────────────
class AnimatedBG(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._t = 0.0
        self._blobs = [
            {"x": 0.15+i*0.11, "y": 0.2+i*0.08, "r": 0.08+i*0.012,
             "dx": 0.0003+i*0.00007, "dy": 0.00025+i*0.00005, "phase": i*0.7}
            for i in range(9)
        ]
        t = QTimer(self)
        t.timeout.connect(self._tick)
        t.start(33)

    def _tick(self):
        self._t += 0.016
        for b in self._blobs:
            b["x"] += b["dx"] * math.sin(self._t * 0.4 + b["phase"])
            b["y"] += b["dy"] * math.cos(self._t * 0.3 + b["phase"])
            b["x"] = max(0.0, min(1.0, b["x"]))
            b["y"] = max(0.0, min(1.0, b["y"]))
        self.update()

    def paintEvent(self, _):
        p = QPainter(self)
        p.setRenderHint(QPainter.RenderHint.Antialiasing)
        w, h = self.width(), self.height()
        wave = (math.sin(self._t * 0.4) + 1) / 2
        grad = QLinearGradient(0, 0, 0, h)
        grad.setColorAt(0, QColor(0, 0, 0))
        grad.setColorAt(wave, QColor(20, 5, 40))
        grad.setColorAt(1, QColor(0x51, 0x0E, 0x8C, 180))
        p.fillRect(0, 0, w, h, grad)
        for i, b in enumerate(self._blobs):
            cx, cy = b["x"]*w, b["y"]*h
            r = b["r"] * min(w, h) * (0.9 + 0.1*math.sin(self._t + b["phase"]))
            rg = QRadialGradient(QPointF(cx, cy), r)
            alpha = int(30 + 15*math.sin(self._t*0.5 + i))
            rg.setColorAt(0, QColor(143, 111, 255, alpha))
            rg.setColorAt(0.5, QColor(81, 14, 140, alpha//2))
            rg.setColorAt(1, QColor(0, 0, 0, 0))
            p.setBrush(rg)
            p.setPen(Qt.PenStyle.NoPen)
            p.drawEllipse(QPointF(cx, cy), r, r)

# ── Hover button ──────────────────────────────────────────────────────────────
_BTN_GREEN      = ("background:#0a2a0a;color:#4ade80;border:1px solid #4ade80;"
                   "border-radius:5px;font-size:11px;padding:4px 8px;")
_BTN_GREEN_H    = ("background:#0d3a0d;color:#6aee90;border:1px solid #6aee90;"
                   "border-radius:5px;font-size:11px;padding:4px 8px;")

class GoldButton(QPushButton):
    def __init__(self, text, parent=None):
        super().__init__(text, parent)
        self._active = False
        self.setStyleSheet(_BTN)
    def setActive(self, on: bool):
        self._active = on
        self.setStyleSheet(_BTN_GREEN if on else _BTN)
    def enterEvent(self, e):
        self.setStyleSheet(_BTN_GREEN_H if self._active else _BTN_H)
        super().enterEvent(e)
    def leaveEvent(self, e):
        self.setStyleSheet(_BTN_GREEN if self._active else _BTN)
        super().leaveEvent(e)


# ── TTS state button ──────────────────────────────────────────────────────────
class TTSStateButton(QPushButton):
    """Play button that reflects TTS state: idle=grey, playing=green, paused=yellow pulse."""
    _IDLE    = ("background:#0d0d12;color:#c8960c;border:1px solid #c8960c;"
                "border-radius:5px;font-size:11px;padding:4px 12px;min-height:24px;")
    _PLAYING = ("background:#0a2a0a;color:#4ade80;border:1px solid #4ade80;"
                "border-radius:5px;font-size:11px;padding:4px 12px;min-height:24px;")

    def __init__(self, text, parent=None):
        super().__init__(text, parent)
        self._state = "idle"
        self._t = 0.0
        self._timer = QTimer(self)
        self._timer.timeout.connect(self._tick)
        self.setStyleSheet(self._IDLE)

    def set_state(self, state: str):
        self._state = state
        self._timer.stop()
        if state == "idle":
            self.setStyleSheet(self._IDLE)
        elif state == "playing":
            self.setStyleSheet(self._PLAYING)
        elif state == "paused":
            self._t = 0.0
            self._timer.start(30)

    def _tick(self):
        FADE = 1.0; HOLD = 0.5; CYCLE = (FADE + HOLD) * 2
        self._t = (self._t + 0.030) % CYCLE
        t = self._t
        if t < FADE:
            b = 1.0 - t / FADE           # fade out bright→dim
        elif t < FADE + HOLD:
            b = 0.0                       # hold dim
        elif t < FADE * 2 + HOLD:
            b = (t - FADE - HOLD) / FADE # fade in dim→bright
        else:
            b = 1.0                       # hold bright
        r = int(0x55 + (0xe8 - 0x55) * b)
        g = int(0x3a + (0xb8 - 0x3a) * b)
        bv = int(0x08 + (0x30 - 0x08) * b)
        col = f"#{r:02x}{g:02x}{bv:02x}"
        self.setStyleSheet(
            f"background:#1a1200;color:{col};border:1px solid {col};"
            "border-radius:5px;font-size:11px;padding:4px 12px;min-height:24px;"
        )

# ── Glass panel ───────────────────────────────────────────────────────────────
class GlassPanel(QFrame):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setStyleSheet(
            "GlassPanel{background:rgba(15,15,18,204);"
            "border:1px solid rgba(143,111,255,80);border-radius:8px;}"
        )

# ── Main window ───────────────────────────────────────────────────────────────
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("DyslexEye")
        self.setMinimumSize(820, 560)
        self.resize(1080, 760)
        self.setAcceptDrops(True)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Window)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self._drag_pos = None

        root = QWidget(self)
        root.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setCentralWidget(root)
        root_layout = QVBoxLayout(root)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)

        self._bg = AnimatedBG(root)
        root_layout.addWidget(self._bg)

        overlay = QWidget(self._bg)
        overlay.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        overlay.setGeometry(self._bg.rect())
        self._bg.resizeEvent = lambda e: overlay.setGeometry(self._bg.rect())
        ov = QVBoxLayout(overlay)
        ov.setContentsMargins(10, 8, 10, 8)
        ov.setSpacing(6)

        # ── Title bar ─────────────────────────────────────────────────────────
        tb = QWidget()
        tb.setStyleSheet("background:transparent;")
        tb.setFixedHeight(36)
        tbl = QHBoxLayout(tb)
        tbl.setContentsMargins(8, 0, 8, 0)
        tbl.setSpacing(8)

        title = QLabel("◎  DyslexEye")
        title.setStyleSheet("color:#c8960c;font-size:13px;font-weight:bold;background:transparent;")
        tbl.addWidget(title)
        tbl.addStretch()

        open_btn = GoldButton("❐ Open File")
        open_btn.setToolTip("Open File — Ctrl+O")
        open_btn.clicked.connect(self._open_dialog)
        tbl.addWidget(open_btn)

        overlay_btn = GoldButton("⌕ Overlay")
        overlay_btn.setToolTip("Overlay mode — coming soon")
        def _overlay_unavailable():
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.information(self, "Not Available",
                "Overlay mode is not yet available.")
        overlay_btn.clicked.connect(_overlay_unavailable)
        tbl.addWidget(overlay_btn)

        tbl.addWidget(self._sep())

        fs_lbl = QLabel("Size")
        fs_lbl.setStyleSheet(_LH)
        tbl.addWidget(fs_lbl)
        self._fs = QSlider(Qt.Orientation.Horizontal)
        self._fs.setRange(12, 36)
        self._fs.setValue(17)
        self._fs.setFixedWidth(80)
        self._fs.setStyleSheet(
            "QSlider::groove:horizontal{height:3px;background:rgba(200,150,12,0.2);border-radius:2px;}"
            "QSlider::handle:horizontal{background:#c8960c;width:12px;height:12px;margin:-5px 0;border-radius:6px;}"
            "QSlider::sub-page:horizontal{background:#c8960c;border-radius:2px;}"
        )
        self._fs.valueChanged.connect(self._set_fontsize)
        tbl.addWidget(self._fs)

        bg_lbl = QLabel("Theme")
        bg_lbl.setStyleSheet(_LH)
        tbl.addWidget(bg_lbl)
        self._bg_pick = QComboBox()
        self._bg_pick.addItems(["Dark gold","Cream","Sky blue","Mint","Lavender","Warm yellow","Navy"])
        self._bg_pick.setStyleSheet(_COMBO)
        self._bg_pick.setFixedWidth(110)
        self._bg_pick.currentIndexChanged.connect(self._set_theme)
        tbl.addWidget(self._bg_pick)

        tbl.addWidget(self._sep())

        self._bionic_on = False
        bionic_btn = GoldButton("½ Half-Bold")
        bionic_btn.setToolTip("Toggle Half-Bold: bolds first half of each word")
        def _toggle_bionic():
            self._bionic_on = not self._bionic_on
            bionic_btn.setActive(self._bionic_on)
            bionic_btn.setText('½ 𝗛𝗮𝗹𝗳-Bold' if self._bionic_on else '½ Half-Bold')
            self._page.runJavaScript("toggleBionic();")
        bionic_btn.clicked.connect(_toggle_bionic)
        tbl.addWidget(bionic_btn)

        self._ruler_on = False
        ruler_btn = GoldButton("▬ Ruler")
        ruler_btn.setToolTip("Toggle Reading Ruler")
        def _toggle_ruler():
            self._ruler_on = not self._ruler_on
            ruler_btn.setActive(self._ruler_on)
            self._page.runJavaScript("toggleRuler(true);" if self._ruler_on else "toggleRuler(false);")
        ruler_btn.clicked.connect(_toggle_ruler)
        tbl.addWidget(ruler_btn)

        tts_play_btn = TTSStateButton("\u25b6 Read")
        tts_play_btn.setToolTip("Start Text-to-Speech")
        def _tts_play():
            self._tts_paused = False
            tts_pause_btn.setText("\u2016 Pause")
            tts_play_btn.set_state("playing")
            self._page.runJavaScript("ttsPlay();")
        tts_play_btn.clicked.connect(_tts_play)
        tbl.addWidget(tts_play_btn)

        self._tts_paused = False
        tts_pause_btn = GoldButton("\u2016 Pause")
        tts_pause_btn.setToolTip("Pause / Resume TTS")
        def _toggle_pause():
            self._tts_paused = not self._tts_paused
            tts_pause_btn.setText("\u25b6 Resume" if self._tts_paused else "\u2016 Pause")
            tts_play_btn.set_state("paused" if self._tts_paused else "playing")
            self._page.runJavaScript("ttsPause();")
        tts_pause_btn.clicked.connect(_toggle_pause)
        tbl.addWidget(tts_pause_btn)

        tts_stop_btn = GoldButton("\u25a0 Stop")
        tts_stop_btn.setToolTip("Stop TTS")
        def _tts_stop():
            self._tts_paused = False
            tts_pause_btn.setText("\u2016 Pause")
            tts_play_btn.set_state("idle")
            self._page.runJavaScript("ttsStop();")
        tts_stop_btn.clicked.connect(_tts_stop)
        tbl.addWidget(tts_stop_btn)

        tbl.addWidget(self._sep())

        min_btn = QPushButton("_")
        min_btn.setFixedSize(26, 26)
        min_btn.setToolTip("Minimise")
        min_btn.setStyleSheet(
            "background:#0d0d12;color:#c8960c;border:1px solid #c8960c;"
            "border-radius:5px;font-size:13px;font-weight:bold;padding-bottom:4px;"
        )
        min_btn.clicked.connect(self.showMinimized)
        tbl.addWidget(min_btn)

        close_btn = QPushButton("✕")
        close_btn.setFixedSize(26, 26)
        close_btn.setToolTip("Close")
        close_btn.setStyleSheet(
            "background:#1a0808;color:#f87171;border:1px solid #f87171;"
            "border-radius:5px;font-size:11px;"
        )
        close_btn.clicked.connect(self.close)
        tbl.addWidget(close_btn)
        ov.addWidget(tb)

        url_row = QWidget()
        url_row.setStyleSheet('background:transparent;')
        url_row.setFixedHeight(32)
        url_rl = QHBoxLayout(url_row)
        url_rl.setContentsMargins(8, 2, 8, 2)
        url_rl.setSpacing(6)
        url_lbl = QLabel('\u2295 URL')
        url_lbl.setStyleSheet(_LH)
        url_rl.addWidget(url_lbl)
        self._url_entry = QLineEdit()
        self._url_entry.setPlaceholderText('https://example.com  — paste URL and press Enter')
        self._url_entry.setStyleSheet(_COMBO)
        self._url_entry.returnPressed.connect(self._load_url_layout)
        url_rl.addWidget(self._url_entry, 1)
        url_go_btn = GoldButton('Load')
        url_go_btn.clicked.connect(self._load_url_layout)
        url_rl.addWidget(url_go_btn)
        self._back_btn = GoldButton('\u2190 Text')
        self._back_btn.setToolTip('Return to text reader mode')
        self._back_btn.clicked.connect(self._exit_site_mode)
        self._back_btn.setVisible(False)
        url_rl.addWidget(self._back_btn)
        ov.addWidget(url_row)

        # ── Glass panel + webview ─────────────────────────────────────────────
        glass = GlassPanel()
        gl = QVBoxLayout(glass)
        gl.setContentsMargins(4, 4, 4, 4)

        self._view = QWebEngineView()
        self._view.setAcceptDrops(False)
        self._view.setStyleSheet("background:transparent;")
        self._page = self._view.page()
        self._page.setBackgroundColor(QColor(0, 0, 0, 0))
        s = self._page.settings()
        s.setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, True)
        s.setAttribute(QWebEngineSettings.WebAttribute.JavascriptEnabled, True)

        self._channel = QWebChannel(self._page)
        self._bridge = Bridge(self)
        self._channel.registerObject("bridge", self._bridge)
        self._page.setWebChannel(self._channel)
        self._page.setHtml(HTML_CONTENT, QUrl("https://localhost/"))

        gl.addWidget(self._view)
        ov.addWidget(glass, 1)
        self._site_mode = False
        self._layout_url = ''
        self._view.loadFinished.connect(self._on_page_loaded)

        self._status = QLabel("Ready — open a file or drag one in")
        self._status.setStyleSheet(_LH + "padding:2px 6px;")
        ov.addWidget(self._status)

        # Shortcuts
        for key, fn in [
            (QKeySequence.StandardKey.Open,  self._open_dialog),
            (QKeySequence.StandardKey.Paste, self._paste),
        ]:
            a = QAction(self)
            a.setShortcut(key)
            a.triggered.connect(fn)
            self.addAction(a)

        self._themes = [
            ("transparent",  "#c8960c", "#e8b830", "transparent"),
            ("#fdf6e3",      "#2a1a00", "#8B6914", "#fdf6e3"),
            ("#e8f4f8",      "#1a3040", "#2980b9", "#e8f4f8"),
            ("#e8f5e8",      "#1a3020", "#27ae60", "#e8f5e8"),
            ("#f5f0ff",      "#2a1a40", "#7b5ea7", "#f5f0ff"),
            ("#fff8e7",      "#3a2800", "#d4890a", "#fff8e7"),
            ("#0d2137",      "#e8e8e8", "#7eb3d4", "#0d2137"),
        ]

    def _on_page_loaded(self, ok: bool):
        if not self._site_mode:
            return
        self._url_entry.setEnabled(True)
        self._url_entry.setPlaceholderText('https://example.com  \u2014 paste URL and press Enter')
        if ok:
            log.info("Layout: direct load OK, injecting font")
            self._inject_dyslexic_font()
        else:
            log.warning("Layout: direct load failed, falling back to HTML fetch")
            self._layout_fallback()

    def _inject_dyslexic_font(self):
        import json as _jf
        css = (
            "@font-face{font-family:'OpenDyslexic';"
            f"src:url('data:font/ttf;base64,{_FONT_B64}')format('truetype')}}" 
            "*,*::before,*::after{font-family:'OpenDyslexic',sans-serif!important;"
            "line-height:1.85!important;letter-spacing:0.05em!important;word-spacing:0.1em!important;}"
        )
        js = ("(function(){"
              "var e=document.getElementById('_de_font');"
              "if(!e){e=document.createElement('style');e.id='_de_font';document.head.appendChild(e);}"
              f"e.textContent={_jf.dumps(css)};"
              "})();")
        self._page.runJavaScript(js)

    def _layout_fallback(self):
        """Option 2: fetch raw HTML, inject font, display via setHtml."""
        import re as _re2, json as _jf2
        log.info(f"Layout fallback: fetching HTML for {self._layout_url}")
        w = HTMLFetchWorker(self._layout_url)
        self._html_fetch_worker = w
        url_snap = self._layout_url
        def _on_fetched(raw_html):
            font_css = (
                "<style id='_de_font'>"
                "@font-face{font-family:'OpenDyslexic';"
                f"src:url('data:font/ttf;base64,{_FONT_B64}')format('truetype')}}" 
                "*,*::before,*::after{font-family:'OpenDyslexic',sans-serif!important;"
                "line-height:1.85!important;letter-spacing:0.05em!important;}"
                "</style>"
                f"<base href='{url_snap}'>"
            )
            if _re2.search(r'<head[^>]*>', raw_html, _re2.IGNORECASE):
                html = _re2.sub(r'(<head[^>]*>)', r'\1' + font_css, raw_html, count=1, flags=_re2.IGNORECASE)
            else:
                html = font_css + raw_html
            self._page.setHtml(html, QUrl(url_snap))
        w.done.connect(_on_fetched)
        w.start()

    def _load_url_layout(self):
        url = self._url_entry.text().strip()
        if not url: return
        if not url.startswith('http'): url = 'https://' + url
        log.info(f"Layout mode: {url}")
        self._site_mode = True
        self._layout_url = url
        self._back_btn.setVisible(True)
        self._url_entry.setEnabled(False)
        self._url_entry.setPlaceholderText('Loading layout\u2026')
        self._view.load(QUrl(url))

    def _exit_site_mode(self):
        log.info("Exiting site mode")
        self._site_mode = False
        self._back_btn.setVisible(False)
        self._url_entry.setEnabled(True)
        self._url_entry.setPlaceholderText('https://example.com  \u2014 paste URL and press Enter')
        self._page.setHtml(HTML_CONTENT, QUrl("https://localhost/"))

    @staticmethod
    def _sep():
        s = QLabel("|")
        s.setStyleSheet("color:#9a7218;background:transparent;font-size:14px;")
        return s

    def _load_url(self):
        url = self._url_entry.text().strip()
        if not url: return
        if not url.startswith('http'): url = 'https://' + url
        log.info(f'Loading URL: {url}')
        self._url_entry.setEnabled(False)
        self._url_entry.setPlaceholderText('Loading…')
        self._page.runJavaScript("show();setReaderText('<p>Loading\u2026</p>','');")
        def _done(result):
            import json as _j
            try:
                d = _j.loads(result)
                if 'error' in d:
                    text = 'Error loading URL: ' + d['error']
                    html_out = '<p style="color:#f87171">' + text + '</p>'
                else:
                    text = d.get('text', '').strip()
                    if not text:
                        text = 'No readable content found at this URL.'
                    paras = [p for p in text.split('\n\n') if p.strip()]
                    html_out = ''.join('<p style="margin:0 0 0.8em 0">' + p.replace('\n', ' ') + '</p>' for p in paras)
                self._page.runJavaScript(
                    'show();' + f'setReaderText({_j.dumps(html_out)},{_j.dumps(text)});'
                )
            except Exception as e:
                log.error(f'_done error: {e}')
            finally:
                self._url_entry.setEnabled(True)
                self._url_entry.setPlaceholderText('https://example.com  — paste URL and press Enter')
                self._page.runJavaScript('show();setReaderText("<p style=\\"color:#f87171\\"">Load error: ' + str(e).replace('"','') + '</p>","");')
        w = URLWorker(url)
        self._url_done_fn = _done  # keep strong ref
        w.done.connect(_done)
        w.start()
        self._url_worker = w

    def _launch_overlay(self):


        if not hasattr(self, "_overlay"):
            self._overlay = OverlayWindow(self)
        self._overlay.show()
        self._overlay.raise_()

    def _open_dialog(self):
        log.info("Open dialog")
        path, _ = QFileDialog.getOpenFileName(self, "Open file", str(Path.home()), SUPPORTED)
        if path:
            log.info(f"Selected: {path}")
            self._load(path)
        else:
            log.debug("Cancelled")

    def _paste(self):
        log.info("Paste triggered")
        try:
            cb = QApplication.clipboard()
            img = cb.image()
            if not img.isNull():
                log.info(f"Clipboard image: {img.width()}x{img.height()}")
                ba = QByteArray()
                buf = QBuffer(ba)
                buf.open(QIODeviceBase.OpenModeFlag.WriteOnly)
                ok = img.save(buf, "PNG")
                buf.close()
                if not ok:
                    log.error("Failed to encode clipboard image")
                    return
                b64 = base64.b64encode(bytes(ba)).decode()
                payload = json.dumps({"name": "pasted_image.png", "ext": "png", "b64": b64})
                self._send(payload)
                self._status.setText("Pasted image")
                return
            text = cb.text()
            if text:
                log.info(f"Clipboard text: {len(text)} chars")
                b64 = base64.b64encode(text.encode()).decode()
                payload = json.dumps({"name": "pasted_text.txt", "ext": "txt", "b64": b64})
                self._send(payload)
                self._status.setText("Pasted text")
            else:
                log.warning("Clipboard empty")
        except Exception as e:
            log.error(f"Paste failed: {e}\n{traceback.format_exc()}")

    def _load(self, path: str):
        log.info(f"Loading: {path}")
        name = Path(path).name
        self._status.setText(f"Loading {name}…")
        self._url_entry.clear()
        loader = FileLoader(path)
        self._file_loader = loader
        def _on_loaded(payload):
            self._send(payload)
            self._status.setText(name)
        loader.done.connect(_on_loaded)
        loader.start()

    def _send(self, payload: str):
        import json as _jj
        self._page.runJavaScript(f"receiveFileData({_jj.dumps(payload)});")

    def _set_fontsize(self, v):
        self._page.runJavaScript(f"document.getElementById('reader').style.fontSize='{v}px';")

    def _set_theme(self, i):
        bg, fg, accent, reader_bg = self._themes[i]
        # null for image theme bg when using default (dark gold)
        img_bg = 'null' if reader_bg == 'transparent' else f"'{reader_bg}'"
        img_fg = 'null' if reader_bg == 'transparent' else f"'{fg}'"
        js = (
            f"(function(){{"
            f"var r=document.getElementById(\'reader\');"
            f"r.style.background=\'{reader_bg}\';"
            f"r.style.color=\'{fg}\';"
            f"document.querySelectorAll(\'#reader h1,#reader h2,#reader h3,#reader h4,#reader h5,#reader h6\').forEach(function(el){{el.style.color=\'{accent}\';}});"
            f"document.querySelectorAll(\'#reader a\').forEach(function(el){{el.style.color=\'{accent}\';}});"
            f"if(window.setThemeBg)window.setThemeBg({img_bg},{img_fg});"
            f"}})();"
        )
        self._page.runJavaScript(js)

    def dragEnterEvent(self, e: QDragEnterEvent):
        if e.mimeData().hasUrls(): e.acceptProposedAction()

    def dropEvent(self, e: QDropEvent):
        for url in e.mimeData().urls():
            p = url.toLocalFile()
            if p:
                log.info(f"Dropped: {p}")
                self._load(p)
                break

    def mousePressEvent(self, e):
        if e.button() == Qt.MouseButton.LeftButton:
            self._drag_pos = e.globalPosition().toPoint() - self.frameGeometry().topLeft()

    def mouseMoveEvent(self, e):
        if self._drag_pos and e.buttons() == Qt.MouseButton.LeftButton:
            self.move(e.globalPosition().toPoint() - self._drag_pos)

    def mouseReleaseEvent(self, e):
        self._drag_pos = None


def main():
    log.info("=== DyslexEye starting ===")
    log.info(f"Python {sys.version}")
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    pal = app.palette()
    pal.setColor(QPalette.ColorRole.Window,     QColor(10, 10, 14))
    pal.setColor(QPalette.ColorRole.WindowText, QColor(200, 150, 12))
    pal.setColor(QPalette.ColorRole.Base,       QColor(10, 10, 14))
    pal.setColor(QPalette.ColorRole.Text,       QColor(200, 150, 12))
    pal.setColor(QPalette.ColorRole.Button,     QColor(13, 13, 18))
    pal.setColor(QPalette.ColorRole.ButtonText, QColor(200, 150, 12))
    app.setPalette(pal)
    log.info("QApplication created")
    global _FONT_FAMILY
    if _FONT_PATH.exists():
        fid = QFontDatabase.addApplicationFont(str(_FONT_PATH))
        if fid >= 0:
            _FONT_FAMILY = QFontDatabase.applicationFontFamilies(fid)[0]
            log.info(f"QFontDatabase: {_FONT_FAMILY}")
    log.info("Creating main window")
    win = MainWindow()
    win.show()
    log.info("Window shown")
    code = app.exec()
    log.info(f"Exit code {code}")
    sys.exit(code)


if __name__ == "__main__":
    main()
